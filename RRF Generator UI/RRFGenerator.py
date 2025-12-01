# RRFGenerator.py - FIXED VERSION

import os
import base64
import datetime
import uuid
from flask import Flask, render_template, request, redirect, url_for, flash, make_response, send_file
import requests
from docx import Document
from datetime import datetime, date
import re
from docx.shared import Pt
import logging
import shutil
from pymongo import MongoClient
import io

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Configuration
UPLOAD_FOLDER = 'Uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
N8N_WEBHOOK_URL = 'http://sara-tech.info/webhook/acb4a2e4-3043-47c8-86d9-2d30d9a06382'

# MongoDB configuration
MONGO_CONFIG = {
    'host': 'localhost',
    'port': 27017,
    'database': 'rrf_processing_db'
}

# Create Flask app
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['MONGO_URI'] = f"mongodb://{MONGO_CONFIG['host']}:{MONGO_CONFIG['port']}/"
app.config['MONGO_DB_NAME'] = MONGO_CONFIG['database']
app.secret_key = os.environ.get('FLASK_SECRET', 'supersecret')

# ============================================
# DATABASE SETUP
# ============================================

# Simple database manager class
class SimpleDatabaseManager:
    def __init__(self):
        self.client = None
        self.db = None
    
    def connect(self):
        try:
            self.client = MongoClient(MONGO_CONFIG['host'], MONGO_CONFIG['port'])
            self.db = self.client[MONGO_CONFIG['database']]
            logger.info("Connected to MongoDB.")
            return self.db
        except Exception as err:
            logger.error(f"MongoDB connection failed: {err}")
            raise
    
    def get_record_by_job_id(self, job_id):
        if self.db is None:
            self.connect()
        collection = self.db.processing_records
        return collection.find_one({'job_id': job_id})
    
    def get_recent_records(self, limit=10, status=None):
        if self.db is None:
            self.connect()
        collection = self.db.processing_records
        query = {}
        if status:
            query['status'] = status
        return list(collection.find(query).sort('run_timestamp', -1).limit(limit))
    
    def get_dashboard_statistics(self):
        if self.db is None:
            self.connect()
        collection = self.db.processing_records
        
        total_success = collection.count_documents({'status': 'SUCCESS'})
        total_documents = collection.count_documents({})
        success_rate = round((total_success / total_documents * 100) if total_documents > 0 else 0, 1)
        
        pipeline = [
            {'$match': {'status': 'SUCCESS'}},
            {'$group': {'_id': '$input_word_name'}}
        ]
        active_templates = len(list(collection.aggregate(pipeline)))
        if active_templates == 0:
            active_templates = 1
        
        return {
            'total_documents': total_success,
            'active_templates': active_templates,
            'avg_processing_time': 2.3,
            'success_rate': success_rate,
            'total_jobs': total_documents
        }

# Initialize database manager
db_manager = SimpleDatabaseManager()
try:
    db_manager.connect()
except Exception as e:
    logger.error(f"Failed to connect to MongoDB: {e}")

# ============================================
# IMPORT AND REGISTER AGENT API
# ============================================

try:
    from agent_api import rrf_agent_api, init_agent_api
    
    # Initialize agent API with database
    init_agent_api(db_manager)
    
    # Register agent API blueprint
    app.register_blueprint(rrf_agent_api, url_prefix='/agent/rrf')
    
    logger.info("✅ Agent API registered successfully!")
except ImportError as e:
    logger.warning(f"⚠️ agent_api.py not found - API endpoints won't be available: {e}")
except Exception as e:
    logger.error(f"❌ Failed to register agent API: {e}")

# ============================================
# DATABASE HELPER FUNCTIONS
# ============================================

def connect_to_db():
    """Connect to MongoDB database."""
    try:
        client = MongoClient(MONGO_CONFIG['host'], MONGO_CONFIG['port'])
        db = client[MONGO_CONFIG['database']]
        logger.info("Connected to MongoDB.")
        return db
    except Exception as err:
        logger.error(f"MongoDB connection failed: {err}")
        raise

def get_dashboard_statistics(db):
    """Get dashboard statistics from MongoDB."""
    collection = db.processing_records
    
    total_documents = collection.count_documents({'status': 'SUCCESS'})
    
    templates_pipeline = [
        {'$match': {'status': 'SUCCESS'}},
        {'$group': {'_id': '$input_word_name'}}
    ]
    active_templates = len(list(collection.aggregate(templates_pipeline)))
    if active_templates == 0:
        active_templates = 1
    
    avg_processing_time = 2.3
    
    total_jobs = collection.count_documents({})
    success_rate = round((total_documents / total_jobs * 100) if total_jobs > 0 else 0, 1)
    
    return {
        'total_documents': total_documents,
        'active_templates': active_templates,
        'avg_processing_time': avg_processing_time,
        'success_rate': success_rate
    }

def save_record(db, job_id, input_pdf_name, input_pdf_path, input_pdf_content,
                input_word_name, input_word_path, input_word_content,
                output_word_name, output_word_path, output_word_content,
                approval_date, log_message, status):
    """Save processing record to MongoDB."""
    collection = db.processing_records
    record = {
        'job_id': job_id,
        'run_timestamp': datetime.now(),
        'input_pdf_name': input_pdf_name,
        'input_pdf_path': input_pdf_path,
        'input_pdf_content': base64.b64encode(input_pdf_content).decode('utf-8') if input_pdf_content else None,
        'input_word_name': input_word_name,
        'input_word_path': input_word_path,
        'input_word_content': base64.b64encode(input_word_content).decode('utf-8') if input_word_content else None,
        'output_word_name': output_word_name,
        'output_word_path': output_word_path,
        'output_word_content': base64.b64encode(output_word_content).decode('utf-8') if output_word_content else None,
        'approval_date': approval_date,
        'log_message': log_message,
        'status': status
    }
    try:
        collection.insert_one(record)
        logger.info(f"Record saved for job_id: {job_id}")
    except Exception as err:
        logger.error(f"Failed to save record to MongoDB: {err}")

# ============================================
# FILE PROCESSING HELPERS
# ============================================

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def send_to_n8n(pdf_b64):
    payload = {'pdf_base64': pdf_b64}
    try:
        resp = requests.post(N8N_WEBHOOK_URL, json=payload, timeout=2000)
        resp.raise_for_status()
        return resp.json()
    except requests.RequestException as e:
        logger.error(f"n8n webhook error: {str(e)}")
        raise

def fill_template(template_path, data, final_output_path):
    doc = Document(template_path)
    formatted_date = datetime.now().strftime('%B %d, %Y')
    approval_date = data.get("approval_date", "")

    placeholders = {
        '<<Date>>': formatted_date,
        '<<Signing Person>>': data['printed_name'],
        '<<Title>>': data['title'],
        '<<Client Name>>': data['org_name'],
        '<<Address>>': f"{data['address']}",
        '<<First Name>>': data['printed_name'].split()[0],
        '<<Fee>>': f"${data['fee']:.2f}",
        '<<Fiscal Year>>': data['fiscal_year'],
        '<<Date1>>': approval_date,
    }

    def replace_placeholder(run_list, placeholder, value):
        full_text = ''.join(run.text for run in run_list)
        if placeholder not in full_text:
            return False
        new_text = full_text.replace(placeholder, value)
        for run in run_list:
            run.text = ''
        if run_list:
            run_list[0].text = new_text
        return True

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            runs = para.runs
            for placeholder, value in placeholders.items():
                if replace_placeholder(runs, placeholder, value):
                    logger.info(f"Replaced {placeholder} with {value}")

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    apply_book_antiqua_font(doc, formatted_date)
    doc.save(final_output_path)
    logger.info(f"Document saved to {final_output_path}")

def apply_book_antiqua_font(doc, approval_date):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Book Antiqua'
        if paragraph.text.startswith("Note:"):
            for run in paragraph.runs:
                run.font.bold = True
        if paragraph.text.startswith("Forms 990 and 199 have been filed"):
            match = re.search(r'([A-Za-z]+ \d{1,2}, \d{4})', paragraph.text)
            if match:
                date_value = match.group(0)
                if date_value:
                    before_date = paragraph.text.split(date_value)[0]
                    after_date = paragraph.text.split(date_value)[1]
                    paragraph.clear()
                    run_before_date = paragraph.add_run(before_date)
                    run_before_date.font.name = 'Book Antiqua'
                    run_before_date.font.size = Pt(12)
                    run_date = paragraph.add_run(date_value)
                    run_date.font.name = 'Book Antiqua'
                    run_date.font.size = Pt(12)
                    run_date.font.bold = True
                    run_after_date = paragraph.add_run(after_date)
                    run_after_date.font.name = 'Book Antiqua'
                    run_after_date.font.size = Pt(12)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Book Antiqua'
                    if paragraph.text.startswith("Forms 990 and 199 have been filed"):
                        match = re.search(r'([A-Za-z]+ \d{1,2}, \d{4})', paragraph.text)
                        if match:
                            date_value = match.group(0)
                            if date_value:
                                before_date = paragraph.text.split(date_value)[0]
                                after_date = paragraph.text.split(date_value)[1]
                                paragraph.clear()
                                run_before_date = paragraph.add_run(before_date)
                                run_before_date.font.name = 'Book Antiqua'
                                run_before_date.font.size = Pt(12)
                                run_date = paragraph.add_run(date_value)
                                run_date.font.name = 'Book Antiqua'
                                run_date.font.size = Pt(12)
                                run_date.font.bold = True
                                run_after_date = paragraph.add_run(after_date)
                                run_after_date.font.name = 'Book Antiqua'
                                run_after_date.font.size = Pt(12)

def format_approval_date(date_str):
    try:
        formatted_date = datetime.strptime(date_str, '%Y-%m-%d').strftime('%B %d, %Y')
        return formatted_date
    except ValueError:
        return ""

def calculate_fee_based_on_revenue(total_revenue):
    try:
        total_revenue = float(total_revenue)
    except ValueError:
        raise ValueError("Invalid total revenue value")
    
    if total_revenue < 50000:
        return 25
    elif 50000 <= total_revenue < 100000:
        return 50
    elif 100001 <= total_revenue < 250000:
        return 75
    elif 250001 <= total_revenue < 1000000:
        return 100
    elif 1000001 <= total_revenue < 5000000:
        return 200
    elif 5000001 <= total_revenue < 20000000:
        return 400
    elif 20000001 <= total_revenue < 100000000:
        return 800
    elif 100000001 <= total_revenue < 500000000:
        return 1000
    else:
        return 1200

# ============================================
# FLASK ROUTES
# ============================================

@app.route('/', methods=['GET'])
def index():
    try:
        db = connect_to_db()
        stats = get_dashboard_statistics(db)
    except Exception as e:
        logger.error(f"Error fetching dashboard statistics: {e}")
        stats = {
            'total_documents': 0,
            'active_templates': 0,
            'avg_processing_time': 0,
            'success_rate': 0
        }
    
    return render_template('index.html', stats=stats)

@app.route('/upload', methods=['POST'])
def upload_files():
    db = connect_to_db()
    
    pdf = request.files.get('pdf_file')
    word = request.files.get('word_file')
    approval_date = request.form.get('approval_date', '').strip()
    
    if not pdf or not word:
        flash('Both files are required.', 'danger')
        save_record(db, '', '', '', None, '', '', None, None, None, None, approval_date, 'Missing input files', 'FAILURE')
        return redirect(url_for('index'))
    
    if not allowed_file(pdf.filename) or not allowed_file(word.filename):
        flash('Only PDF and DOCX files are allowed.', 'danger')
        save_record(db, '', pdf.filename, '', None, word.filename, '', None, None, None, None, approval_date, 'Invalid file extensions', 'FAILURE')
        return redirect(url_for('index'))
    
    formatted_approval_date = format_approval_date(approval_date)
    
    job_id = str(uuid.uuid4())
    job_folder = os.path.join(app.config['UPLOAD_FOLDER'], job_id)
    os.makedirs(job_folder, exist_ok=True)
    
    pdf_path = os.path.join(job_folder, pdf.filename)
    word_path = os.path.join(job_folder, word.filename)
    pdf.save(pdf_path)
    word.save(word_path)
    
    input_pdf_content = None
    input_word_content = None
    try:
        with open(pdf_path, 'rb') as f:
            input_pdf_content = f.read()
        with open(word_path, 'rb') as f:
            input_word_content = f.read()
    except Exception as e:
        logger.error(f"Failed to read input files: {str(e)}")
        save_record(db, job_id, pdf.filename, pdf_path, None, word.filename, word_path, None, None, None, None, approval_date, f"File read error: {str(e)}", 'FAILURE')
        flash(f"File read error: {e}", 'danger')
        return redirect(url_for('index'))
    
    pdf_b64 = base64.b64encode(input_pdf_content).decode()
    
    try:
        result_json = send_to_n8n(pdf_b64)
        
        if isinstance(result_json, list):
            if not result_json:
                raise ValueError("n8n returned an empty list")
            extracted = result_json[0]
        else:
            extracted = result_json
        
        printed_name = extracted.get('Signing Person', '').strip()
        if not printed_name:
            raise ValueError("Missing 'printed_name' in extracted data")
        
        first_name = printed_name.split()[0]
        title = extracted.get('Title', '').strip()
        org_name = extracted.get('Client name', '').strip()
        addr1 = extracted.get('Address Line 1', '').strip()
        addr2 = extracted.get('Address Line 2', '').strip()
        address = addr1 + ("\n" + addr2 if addr2 else "")
        fiscal_year = extracted.get('Fiscal Year', '').strip()
        total_revenue = extracted.get('Total Revenue', '')
        
        fee = calculate_fee_based_on_revenue(total_revenue)
        
        data = {
            'printed_name': printed_name,
            'first_name': first_name,
            'title': title,
            'org_name': org_name,
            'address': address,
            'fiscal_year': fiscal_year,
            'fee': fee,
            'approval_date': formatted_approval_date
        }
        
        pdf_filename_without_ext = os.path.splitext(pdf.filename)[0]
        prefix = pdf_filename_without_ext.split('_')[0]
        final_filename = f"{prefix}_Letter of Instruction- Form 990 & RRF-1.docx"
        temp_output_path = os.path.join(job_folder, "Letter of Instruction- Form 990 & RRF-1.docx")
        fill_template(word_path, data, temp_output_path)
        
        final_output_path = os.path.join(job_folder, final_filename)
        shutil.copy(temp_output_path, final_output_path)
        logger.info(f"Final file saved as: {final_output_path}")
        
        output_word_content = None
        try:
            with open(final_output_path, 'rb') as f:
                output_word_content = f.read()
        except Exception as e:
            logger.error(f"Failed to read output file for DB storage: {str(e)}")
            save_record(db, job_id, pdf.filename, pdf_path, input_pdf_content,
                        word.filename, word_path, input_word_content,
                        final_filename, final_output_path, None,
                        approval_date, f"Output file read error for DB: {str(e)}", 'FAILURE')
            flash(f"Output file read error for DB: {e}", 'danger')
            return redirect(url_for('result', job_id=job_id))
        
        save_record(
            db, job_id, pdf.filename, pdf_path, input_pdf_content,
            word.filename, word_path, input_word_content,
            final_filename, final_output_path, output_word_content,
            formatted_approval_date, f"Processed successfully: {result_json}", 'SUCCESS'
        )
        
        try:
            shutil.rmtree(job_folder)
            logger.info(f"Cleaned up local job folder: {job_folder}")
        except Exception as e:
            logger.warning(f"Failed to clean up local job folder {job_folder}: {e}")
        
        return redirect(url_for('result', job_id=job_id))
    
    except Exception as e:
        logger.error(f"Processing error: {str(e)}")
        save_record(
            db, job_id, pdf.filename, pdf_path, input_pdf_content,
            word.filename, word_path, input_word_content,
            None, None, None, formatted_approval_date, f"Processing error: {str(e)}", 'FAILURE'
        )
        flash(f"Processing error: {e}", 'danger')
        return redirect(url_for('result', job_id=job_id))

@app.route('/result/<job_id>')
def result(job_id):
    db = connect_to_db()
    collection = db.processing_records
    record = collection.find_one({'job_id': job_id})
    
    job = {
        'id': job_id,
        'status': record['status'] if record else 'not_found',
        'output_filename': record['output_word_name'] if record and record['status'] == 'SUCCESS' else None,
        'created_at': record['run_timestamp'] if record else None,
        'error_message': record.get('log_message', '') if record and record['status'] == 'FAILURE' else None
    }
    return render_template('result.html', job=job)

@app.route('/download/<job_id>')
def download_file(job_id):
    db = connect_to_db()
    collection = db.processing_records
    record = collection.find_one({'job_id': job_id})
    
    if record and 'output_word_content' in record and record['output_word_content']:
        output_word_name = record.get('output_word_name', f"{job_id}_output.docx")
        output_word_content_b64 = record['output_word_content']
        
        try:
            output_word_binary = base64.b64decode(output_word_content_b64.encode('utf-8'))
            buffer = io.BytesIO(output_word_binary)
            buffer.seek(0)
            
            response = make_response(buffer.read())
            response.headers["Content-Disposition"] = f"attachment; filename={output_word_name}"
            response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            return response
        except Exception as e:
            logger.error(f"Error serving file from MongoDB for job_id {job_id}: {e}")
            flash(f"Error serving file: {e}", "danger")
            return redirect(url_for('result', job_id=job_id))
    else:
        flash("Completed document not found in database.", "danger")
        return redirect(url_for('result', job_id=job_id))

# ============================================
# RUN APP
# ============================================

if __name__ == '__main__':
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    logger.info("🚀 Starting RRF Generator...")
    logger.info(f"📁 Upload folder: {UPLOAD_FOLDER}")
    logger.info(f"🗄️  MongoDB: {MONGO_CONFIG['host']}:{MONGO_CONFIG['port']}/{MONGO_CONFIG['database']}")
    app.run(host="0.0.0.0", port=5009, debug=False)