"""
Template Filler - Word Document Processing
Fills Word template with extracted data
"""

import logging
from docx import Document
from docx.shared import Pt
from datetime import datetime
import re

logger = logging.getLogger(__name__)


class WordTemplateFiller:
    """Fill Word template with extracted data"""
    
    def __init__(self):
        pass
    
    def fill_template(self, template_path: str, data: dict, output_path: str):
        """
        Fill Word template with data
        
        Args:
            template_path: Path to Word template
            data: Dictionary with extracted data
            output_path: Path to save filled document
        """
        try:
            logger.info(f"📝 Filling template: {template_path}")
            
            doc = Document(template_path)
            
            # Prepare data
            formatted_date = datetime.now().strftime('%B %d, %Y')
            approval_date = data.get("approval_date", formatted_date)
            
            # Build placeholders
            placeholders = {
                '<<Date>>': formatted_date,
                '<<Signing Person>>': data.get('printed_name', ''),
                '<<Title>>': data.get('title', ''),
                '<<Client Name>>': data.get('org_name', ''),
                '<<Address>>': data.get('address', ''),
                '<<First Name>>': data.get('first_name', ''),
                '<<Fee>>': f"${data.get('fee', 0):.2f}",
                '<<Fiscal Year>>': data.get('fiscal_year', ''),
                '<<Date1>>': approval_date,
            }
            
            logger.info(f"Replacing {len(placeholders)} placeholders")
            
            # Replace in paragraphs
            self._process_paragraphs(doc.paragraphs, placeholders)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._process_paragraphs(cell.paragraphs, placeholders)
            
            # Apply Book Antiqua font
            self._apply_font_styling(doc, approval_date)
            
            # Save document
            doc.save(output_path)
            logger.info(f"✅ Document saved to: {output_path}")
            
        except Exception as e:
            logger.error(f"❌ Template filling error: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to fill template: {str(e)}")
    
    def _process_paragraphs(self, paragraphs, placeholders):
        """Replace placeholders in paragraphs"""
        for para in paragraphs:
            runs = para.runs
            for placeholder, value in placeholders.items():
                if self._replace_placeholder(runs, placeholder, value):
                    logger.debug(f"Replaced {placeholder} with {value}")
    
    def _replace_placeholder(self, run_list, placeholder, value):
        """Replace placeholder in run list"""
        full_text = ''.join(run.text for run in run_list)
        
        if placeholder not in full_text:
            return False
        
        new_text = full_text.replace(placeholder, value)
        
        # Clear all runs
        for run in run_list:
            run.text = ''
        
        # Set new text in first run
        if run_list:
            run_list[0].text = new_text
        
        return True
    
    def _apply_font_styling(self, doc, approval_date):
        """Apply Book Antiqua font and special formatting"""
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            # Apply Book Antiqua to all runs
            for run in paragraph.runs:
                run.font.name = 'Book Antiqua'
            
            # Bold "Note:" paragraphs
            if paragraph.text.startswith("Note:"):
                for run in paragraph.runs:
                    run.font.bold = True
            
            # Special formatting for filing date paragraph
            if paragraph.text.startswith("Forms 990 and 199 have been filed"):
                self._format_filing_date_paragraph(paragraph)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Book Antiqua'
                        
                        if paragraph.text.startswith("Forms 990 and 199 have been filed"):
                            self._format_filing_date_paragraph(paragraph)
    
    def _format_filing_date_paragraph(self, paragraph):
        """Format paragraph with filing date (bold the date)"""
        # Extract date from paragraph
        match = re.search(r'([A-Za-z]+ \d{1,2}, \d{4})', paragraph.text)
        
        if match:
            date_value = match.group(0)
            before_date = paragraph.text.split(date_value)[0]
            after_date = paragraph.text.split(date_value)[1]
            
            # Rebuild paragraph with formatted date
            paragraph.clear()
            
            # Before date
            run_before = paragraph.add_run(before_date)
            run_before.font.name = 'Book Antiqua'
            run_before.font.size = Pt(12)
            
            # Date (bold)
            run_date = paragraph.add_run(date_value)
            run_date.font.name = 'Book Antiqua'
            run_date.font.size = Pt(12)
            run_date.font.bold = True
            
            # After date
            run_after = paragraph.add_run(after_date)
            run_after.font.name = 'Book Antiqua'
            run_after.font.size = Pt(12)