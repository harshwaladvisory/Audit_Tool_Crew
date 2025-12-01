import tempfile
import os
import logging
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

class DocumentGenerator:
    """Generate Word and PDF documents from cover letter data"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def generate_docx(self, cover_letter_data: Dict[str, Any], client_name: str) -> str:
        """
        Generate Word document from cover letter data
        
        Args:
            cover_letter_data: Processed cover letter data
            client_name: Name of the client
            
        Returns:
            Path to the generated temporary file
        """
        try:
            # Create new document
            doc = Document()
            
            # Set wider margins (reduce by 0.5 inch from both sides)
            sections = doc.sections
            for section in sections:
                section.left_margin = Inches(1.0)   # Default is 1.25, reduce by 0.5 
                section.right_margin = Inches(1.0)  # Default is 1.25, reduce by 0.5
            
            # Set document style - Palatino Linotype, 11pt, Dark Blue
            from docx.shared import Pt
            
            # Add introduction line with draft type and tax year
            intro_text = f"We bring the following to your attention regarding the preparation of the {cover_letter_data['draft_type'].lower()} of {cover_letter_data['tax_year']} tax returns of {client_name}:"
            intro_para = doc.add_paragraph(intro_text)
            intro_para.paragraph_format.space_after = Pt(12)  # Single line spacing after intro
            intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify text
            
            # Set intro text formatting
            for run in intro_para.runs:
                run.font.name = 'Palatino Linotype'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Add sections with proper grouping
            current_section_number = 1
            
            for section in cover_letter_data['sections']:
                # Create main paragraph for this section
                section_para = doc.add_paragraph()
                section_para.paragraph_format.space_after = Pt(12)  # Single line spacing between points
                # No justify alignment for headers - they stay left aligned
                
                # Add numbering
                number_text = f"{current_section_number}. "
                if section.get('custom_numbering'):
                    number_text = f"{section['custom_numbering']} "
                
                number_run = section_para.add_run(number_text)
                number_run.bold = True
                number_run.font.name = 'Palatino Linotype'
                number_run.font.size = Pt(11)
                number_run.font.color.rgb = RGBColor(0, 51, 102)
                
                # Add header in bold
                if section['header']:
                    header_run = section_para.add_run(f"{section['header']}")
                    header_run.bold = True
                    header_run.font.name = 'Palatino Linotype'
                    header_run.font.size = Pt(11)
                    header_run.font.color.rgb = RGBColor(0, 51, 102)
                
                # If there are multiple instructions for this header, add them as bullets
                if 'instructions' in section and len(section['instructions']) > 1:
                    # Add line break after header
                    section_para.add_run('\n')
                    
                    # Add each instruction as a bullet point
                    for i, instruction_item in enumerate(section['instructions']):
                        bullet_letter = chr(ord('a') + i)  # a), b), c), etc.
                        bullet_para = doc.add_paragraph()
                        bullet_para.paragraph_format.space_after = Pt(12)  # Single line spacing between subpoints
                        bullet_para.paragraph_format.left_indent = Inches(0.5)
                        bullet_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify text
                        
                        bullet_run = bullet_para.add_run(f"{bullet_letter}) ")
                        bullet_run.font.name = 'Palatino Linotype'
                        bullet_run.font.size = Pt(11)
                        bullet_run.font.color.rgb = RGBColor(0, 51, 102)
                        
                        instruction_run = bullet_para.add_run(instruction_item['instruction'])
                        instruction_run.font.name = 'Palatino Linotype'
                        instruction_run.font.size = Pt(11)
                        instruction_run.font.color.rgb = RGBColor(0, 51, 102)
                else:
                    # Single instruction, add it directly
                    if section['header']:
                        section_para.add_run(' ')
                    
                    instruction_text = section.get('instruction', section['instructions'][0]['instruction'] if 'instructions' in section else '')
                    instruction_run = section_para.add_run(instruction_text)
                    instruction_run.font.name = 'Palatino Linotype'
                    instruction_run.font.size = Pt(11)
                    instruction_run.font.color.rgb = RGBColor(0, 51, 102)
                    
                    # Apply justify alignment only to instruction text, not headers
                    section_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                current_section_number += 1
            
            # Add closing paragraph
            closing_para = doc.add_paragraph("Please review all the pages of the return including Schedule O and let us know if any changes are required.")
            closing_para.paragraph_format.space_before = Pt(12)
            closing_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Set closing text formatting
            for run in closing_para.runs:
                run.font.name = 'Palatino Linotype'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            self.logger.debug(f"Generated Word document: {temp_file.name}")
            return temp_file.name
            
        except Exception as e:
            self.logger.error(f"Error generating Word document: {str(e)}")
            raise
    
    def generate_pdf(self, cover_letter_data: Dict[str, Any], client_name: str) -> str:
        """
        Generate PDF document from cover letter data
        
        Args:
            cover_letter_data: Processed cover letter data
            client_name: Name of the client
            
        Returns:
            Path to the generated temporary file
        """
        try:
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_file.close()
            
            # Create PDF document with wider margins (reduce by 0.5 inch = 36 points from both sides)
            doc = SimpleDocTemplate(temp_file.name, pagesize=letter,
                                  rightMargin=36, leftMargin=36,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles with Palatino Linotype, 11pt, Dark Blue
            from reportlab.lib import colors
            dark_blue = colors.HexColor('#003366')  # Dark Blue
            
            # Style for headers (no justify alignment)
            header_style = ParagraphStyle(
                'HeaderStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,  # Single line spacing between points
                fontName='Times-Roman',  # Closest to Palatino Linotype
                textColor=dark_blue,
                leftIndent=0
                # No alignment specified - stays left aligned
            )
            
            # Style for instruction text (justified)
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,  # Single line spacing between points
                fontName='Times-Roman',  # Closest to Palatino Linotype
                textColor=dark_blue,
                leftIndent=0,
                alignment=0  # Justify text (0 = justify, 1 = center, 2 = right)
            )
            
            bullet_style = ParagraphStyle(
                'BulletStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,  # Single line spacing between subpoints
                fontName='Times-Roman',
                textColor=dark_blue,
                leftIndent=36,  # Indent for bullet points
                bulletIndent=18,
                alignment=0  # Justify text
            )
            
            # Build content
            content = []
            
            # Add introduction line with draft type and tax year
            intro_text = f"We bring the following to your attention regarding the preparation of the {cover_letter_data['draft_type'].lower()} of {cover_letter_data['tax_year']} tax returns of {client_name}:"
            content.append(Paragraph(intro_text, section_style))
            content.append(Spacer(1, 12))  # Single line spacing after intro
            
            # Add sections
            current_section_number = 1
            
            # Add sections with proper grouping
            for section in cover_letter_data['sections']:
                # Create section text
                number_text = f"{current_section_number}. "
                if section.get('custom_numbering'):
                    number_text = f"{section['custom_numbering']} "
                
                section_text = f"<b>{number_text}"
                
                if section['header']:
                    section_text += f"{section['header']}</b>"
                else:
                    section_text += "</b>"
                
                # If there are multiple instructions for this header, add them as bullets
                if 'instructions' in section and len(section['instructions']) > 1:
                    # Add header (not justified)
                    content.append(Paragraph(section_text, header_style))
                    
                    # Add each instruction as a bullet point (justified)
                    for i, instruction_item in enumerate(section['instructions']):
                        bullet_letter = chr(ord('a') + i)  # a), b), c), etc.
                        bullet_text = f"{bullet_letter}) {instruction_item['instruction']}"
                        content.append(Paragraph(bullet_text, bullet_style))
                else:
                    # Single instruction - combine header and instruction (justified)
                    if section['header']:
                        section_text += " "
                    
                    instruction_text = section.get('instruction', section['instructions'][0]['instruction'] if 'instructions' in section else '')
                    section_text += instruction_text
                    content.append(Paragraph(section_text, section_style))
                
                current_section_number += 1
            
            # Add closing paragraph
            content.append(Spacer(1, 12))
            content.append(Paragraph("Please review all the pages of the return including Schedule O and let us know if any changes are required.", section_style))
            
            # Build PDF
            doc.build(content)
            
            self.logger.debug(f"Generated PDF document: {temp_file.name}")
            return temp_file.name
            
        except Exception as e:
            self.logger.error(f"Error generating PDF document: {str(e)}")
            raise
    
    def cleanup_temp_file(self, filepath: str):
        """
        Clean up temporary file
        
        Args:
            filepath: Path to the temporary file
        """
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                self.logger.debug(f"Cleaned up temporary file: {filepath}")
        except Exception as e:
            self.logger.warning(f"Could not clean up temporary file {filepath}: {str(e)}")
